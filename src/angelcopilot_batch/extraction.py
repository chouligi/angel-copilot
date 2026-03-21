from __future__ import annotations

import io
from pathlib import Path
import zipfile

from angelcopilot_batch.models import EvidenceBlock, EvidenceBundle


def extract_evidence_bundle(file_paths: list[Path]) -> EvidenceBundle:
    bundle = EvidenceBundle()
    seen_text: set[str] = set()

    for file_path in file_paths:
        extension = file_path.suffix.lower()
        if extension == ".zip":
            _extract_zip_archive(file_path=file_path, bundle=bundle, seen_text=seen_text)
            continue

        extractor = _get_extractor(extension)
        if extractor is None:
            bundle.warnings.append(f"Unsupported file type: {file_path}")
            continue

        try:
            text = extractor(file_path).strip()
        except Exception as exc:  # noqa: BLE001
            bundle.warnings.append(f"Failed to parse {file_path}: {exc}")
            continue

        if not text:
            continue

        _append_text_block(bundle=bundle, seen_text=seen_text, source_path=file_path, text=text)

    return bundle


def _extract_zip_archive(file_path: Path, bundle: EvidenceBundle, seen_text: set[str]) -> None:
    try:
        with zipfile.ZipFile(file_path) as archive:
            handled_any_supported = False
            for member in sorted(archive.infolist(), key=lambda item: item.filename):
                if member.is_dir():
                    continue

                member_suffix = Path(member.filename).suffix.lower()
                if _get_extractor(member_suffix) is None:
                    continue

                handled_any_supported = True
                raw_bytes = archive.read(member.filename)
                try:
                    text = _extract_from_bytes(member_suffix, raw_bytes).strip()
                except Exception as exc:  # noqa: BLE001
                    bundle.warnings.append(f"Failed to parse archive member {member.filename}: {exc}")
                    continue

                source = Path(f"{file_path}!{member.filename}")
                _append_text_block(bundle=bundle, seen_text=seen_text, source_path=source, text=text)

            if not handled_any_supported:
                bundle.warnings.append(f"No supported docs found in archive: {file_path}")
    except zipfile.BadZipFile as exc:
        bundle.warnings.append(f"Invalid zip archive {file_path}: {exc}")


def _get_extractor(extension: str):  # noqa: ANN202
    if extension in {".txt", ".md"}:
        return _extract_text_file
    if extension == ".pdf":
        return _extract_pdf_file
    if extension == ".docx":
        return _extract_docx_file
    return None


def _extract_text_file(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf_file(file_path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("Install pypdf to parse PDF files") from exc

    reader = PdfReader(str(file_path))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(page for page in pages if page)


def _extract_docx_file(file_path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("Install python-docx to parse DOCX files") from exc

    document = Document(str(file_path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def _extract_from_bytes(extension: str, payload: bytes) -> str:
    if extension in {".txt", ".md"}:
        return payload.decode("utf-8", errors="ignore")
    if extension == ".pdf":
        return _extract_pdf_bytes(payload)
    if extension == ".docx":
        return _extract_docx_bytes(payload)
    raise RuntimeError(f"Unsupported archive member type: {extension}")


def _extract_pdf_bytes(payload: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("Install pypdf to parse PDF files") from exc

    reader = PdfReader(io.BytesIO(payload))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(page for page in pages if page)


def _extract_docx_bytes(payload: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("Install python-docx to parse DOCX files") from exc

    document = Document(io.BytesIO(payload))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def _append_text_block(bundle: EvidenceBundle, seen_text: set[str], source_path: Path, text: str) -> None:
    if not text:
        return

    text_key = " ".join(text.split())
    if text_key in seen_text:
        return

    seen_text.add(text_key)
    bundle.evidence_blocks.append(EvidenceBlock(source_path=source_path, text=text))
